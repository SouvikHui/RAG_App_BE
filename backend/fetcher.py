import requests
from typing import List
from langchain.docstore.document import Document
from bs4 import BeautifulSoup
from fastapi import UploadFile
from io import BytesIO
import docx  # python-docx
from pypdf import PdfReader
# import mimetypes

def load_urls(urls:List[str]) -> List[Document]:
    docs = []
    for url in urls:
        try:
            response = requests.get(url)
            content_type = response.headers.get("Content-Type", "")

            if "application/pdf" in content_type or url.lower().endswith(".pdf"):
                from pypdf import PdfReader
                from io import BytesIO

                reader = PdfReader(BytesIO(response.content))
                text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
                docs.append(Document(page_content=text, metadata={"source": url}))

            elif "text/plain" in content_type or url.lower().endswith(".txt"):
                text = response.text
                docs.append(Document(page_content=text, metadata={"source": url}))

            else:  # Assume HTML
                soup = BeautifulSoup(response.text, "html.parser")
                text = soup.get_text(separator=" ", strip=True)
                docs.append(Document(page_content=text, metadata={"source": url}))

        except Exception as e:
            print(f"[!] Failed to process URL: {url} - Error: {str(e)}")

    return docs

# Uploaded Files
async def process_uploaded_file(file: UploadFile) -> List[Document]:
    content = await file.read()
    file_type = file.filename.lower()
    docs = []
    if file_type.endswith(".pdf"):
        pdf_reader = PdfReader(BytesIO(content))
        text = "\n".join([p.extract_text() for p in pdf_reader.pages if p.extract_text()])
    elif file_type.endswith(".txt"):
        text = content.decode("utf-8")
    elif file_type.endswith(".docx"):
        doc = docx.Document(BytesIO(content))
        text = "\n".join([p.text for p in doc.paragraphs])
    elif file_type.endswith((".csv",".xls", ".xlsx")):
        csv_reader = BytesIO(content)
        try:
            if file_type.endswith(".csv"):
                df = pd.read_csv(csv_reader)
            elif file_type.endswith((".xls", ".xlsx")):
                df = pd.read_excel(csv_reader)
            text = df.to_string(index=False)
        except Exception as e:
            text = f"Error reading spreadsheet: {str(e)}"
    else:
        raise ValueError("Unsupported file type")
    docs.append(Document(page_content=text, metadata={"source": file.filename}))

    return docs
